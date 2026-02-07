"""
RAG System for Client Document Analysis
Ingests DOCX files and creates embeddings for semantic search
"""
import os
import json
from pathlib import Path
from typing import List, Dict, Any
import chromadb
from chromadb.config import Settings
from docx import Document
from datetime import datetime


class RAGSystem:
    def __init__(self, persist_directory: str = None):
        """Initialize the RAG system with ChromaDB."""
        if persist_directory is None:
            # Default to backend/data/chroma_db
            persist_directory = str(Path(__file__).parent / "data" / "chroma_db")
            
        self.persist_directory = persist_directory
        Path(persist_directory).mkdir(parents=True, exist_ok=True)
        
        self.client = chromadb.PersistentClient(path=persist_directory)
        
        # Create or get collection
        try:
            self.collection = self.client.get_collection(name="client_documents")
        except:
            self.collection = self.client.create_collection(
                name="client_documents",
                metadata={"description": "Client documents and context"}
            )
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from a DOCX file."""
        try:
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return "\n".join(full_text)
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks."""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def ingest_document(self, file_path: str, metadata: Dict[str, Any] = None) -> int:
        """Ingest a single document into the RAG system."""
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return 0
        
        # Extract text
        text = self.extract_text_from_docx(file_path)
        if not text:
            print(f"No text extracted from {file_path}")
            return 0
        
        # Create chunks
        chunks = self.chunk_text(text)
        
        # Prepare metadata
        base_metadata = {
            "source": os.path.basename(file_path),
            "file_path": file_path,
            "ingested_at": datetime.now().isoformat(),
            "chunk_count": len(chunks)
        }
        
        if metadata:
            base_metadata.update(metadata)
        
        # Add to ChromaDB
        ids = []
        metadatas = []
        documents = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = f"{os.path.basename(file_path)}_chunk_{i}"
            chunk_metadata = base_metadata.copy()
            chunk_metadata["chunk_index"] = i
            
            ids.append(chunk_id)
            metadatas.append(chunk_metadata)
            documents.append(chunk)
        
        self.collection.add(
            ids=ids,
            documents=documents,
            metadatas=metadatas
        )
        
        print(f"✅ Ingested {len(chunks)} chunks from {os.path.basename(file_path)}")
        return len(chunks)
    
    def ingest_directory(self, directory_path: str) -> Dict[str, int]:
        """Ingest all DOCX files from a directory."""
        directory = Path(directory_path)
        results = {}
        
        if not directory.exists():
            print(f"Directory not found: {directory_path}")
            return results
        
        docx_files = list(directory.glob("*.docx"))
        
        if not docx_files:
            print(f"No DOCX files found in {directory_path}")
            return results
        
        print(f"\n📚 Ingesting {len(docx_files)} documents...")
        
        for file_path in docx_files:
            # Skip temporary files
            if file_path.name.startswith("~$"):
                continue
            
            chunk_count = self.ingest_document(str(file_path))
            results[file_path.name] = chunk_count
        
        print(f"\n✅ Ingestion complete! Total documents: {len(results)}")
        return results
    
    def search(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        """Search for relevant documents."""
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results
            )
            
            formatted_results = []
            
            if results and results['documents'] and len(results['documents']) > 0:
                for i in range(len(results['documents'][0])):
                    formatted_results.append({
                        "content": results['documents'][0][i],
                        "metadata": results['metadatas'][0][i] if results['metadatas'] else {},
                        "distance": results['distances'][0][i] if results['distances'] else None
                    })
            
            return formatted_results
        except Exception as e:
            print(f"Search error: {e}")
            return []
    
    def get_all_documents(self) -> List[Dict[str, Any]]:
        """Get all documents in the collection."""
        try:
            results = self.collection.get()
            
            documents = []
            if results and results['documents']:
                for i in range(len(results['documents'])):
                    documents.append({
                        "id": results['ids'][i],
                        "content": results['documents'][i],
                        "metadata": results['metadatas'][i] if results['metadatas'] else {}
                    })
            
            return documents
        except Exception as e:
            print(f"Error getting documents: {e}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the RAG system."""
        try:
            count = self.collection.count()
            
            # Get unique sources
            all_docs = self.get_all_documents()
            sources = set()
            for doc in all_docs:
                if 'source' in doc['metadata']:
                    sources.add(doc['metadata']['source'])
            
            return {
                "total_chunks": count,
                "total_documents": len(sources),
                "sources": list(sources)
            }
        except Exception as e:
            print(f"Error getting stats: {e}")
            return {"total_chunks": 0, "total_documents": 0, "sources": []}


# Standalone functions for easy import
def create_rag_system() -> RAGSystem:
    """Create and return a RAG system instance."""
    return RAGSystem()


def ingest_documents(directory_path: str = None) -> Dict[str, int]:
    """Ingest all documents from the specified directory."""
    if directory_path is None:
        directory_path = str(Path(__file__).parent / "data" / "client_documents")
    rag = create_rag_system()
    return rag.ingest_directory(directory_path)


if __name__ == "__main__":
    # Test the RAG system
    print("🚀 Initializing RAG System...")
    rag = RAGSystem()
    
    # Ingest documents
    results = rag.ingest_directory("./data/client_documents")
    
    # Show stats
    stats = rag.get_stats()
    print(f"\n📊 RAG System Stats:")
    print(f"   Total chunks: {stats['total_chunks']}")
    print(f"   Total documents: {stats['total_documents']}")
    print(f"   Sources: {', '.join(stats['sources'])}")
